# content_extractor.py

from docx import Document
import os
import requests
from bs4 import BeautifulSoup

def clean_text(text):
    """
    Cleans the input text by stripping excessive whitespace.

    Args:
        text (str): The raw text to clean.

    Returns:
        str: Cleaned text with no extra whitespace.
    """
    return ' '.join(text.split()).strip()

def extract_content_from_urls(urls, folder_name):
    """
    Extracts content (h1, h2, p, span) from a list of URLs and saves the text into .docx files.

    Args:
        urls (list): A list of URLs to scrape.
        folder_name (str): The folder where .docx files will be saved.

    Returns:
        None
    """
    for url in urls:
        try:
            response = requests.get(url)
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Create a new document for each URL
            document = Document()
            
            # Extract content from specified tags and add to the document
            for tag in ['h1', 'h2', 'p', 'span']:
                elements = soup.find_all(tag)
                for element in elements:
                    cleaned_text = clean_text(element.get_text())
                    if tag in ['h1', 'h2']:  # Format headers in bold
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run(cleaned_text)
                        run.bold = True
                    else:  # Add normal text
                        document.add_paragraph(cleaned_text)
            
            # Create a safe filename based on the URL
            safe_url = url.replace('https://', '').replace('http://', '').replace('/', '_')[:50]
            docx_file_name = os.path.join(folder_name, f"{safe_url}.docx")
            
            # Save the extracted content into a .docx file
            document.save(docx_file_name)
            print(f"Content extracted and saved to {docx_file_name}")
            
        except Exception as e:
            print(f"Failed to fetch {url}: {e}")
