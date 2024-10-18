import ast
import os
from docx import Document
from .doc_generator import generate_markdown  # Import the generate_markdown function
from .code_analyzer import CodeAnalyzer
from .summary_generator import generate_summary

def generate_documentation(output_folder):
    # Specify the path to your_code.py
    script_dir = os.path.dirname(__file__)  # Get the directory of the current script
    code_file_path = os.path.join(script_dir, 'your_code.py')  # Update the filename if needed
    
    # Analyze the code
    with open(code_file_path, 'r') as file:
        tree = ast.parse(file.read())
        analyzer = CodeAnalyzer()
        analyzer.visit(tree)

    # Get the function data for the first function
    function_data = analyzer.function_data[0]
    
    # Generate summaries for different audiences
    dev_summary = generate_summary(function_data['docstring'], audience="developer")
    user_summary = generate_summary(function_data['docstring'], audience="end_user")
    
    # Generate Markdown for the function
    markdown_doc = generate_markdown(function_data)

    # Create a new Document
    doc = Document()
    doc.add_heading('Generated Documentation', level=1)

    # Add summaries to the document
    doc.add_heading('Developer Summary', level=2)
    doc.add_paragraph(dev_summary)

    doc.add_heading('End User Summary', level=2)
    doc.add_paragraph(user_summary)

    # Add the Markdown content to the document (converting to plain text)
    doc.add_heading('Function Documentation', level=2)
    doc.add_paragraph(markdown_doc)  # Adding the Markdown directly

    # Save the document
    os.makedirs(output_folder, exist_ok=True)  # Create the folder if it doesn't exist
    doc_path = os.path.join(output_folder, 'documentation.docx')
    doc.save(doc_path)

    print(f"Documentation saved at: {doc_path}")

# Call the function and specify the output folder
if __name__ == "__main__":
    generate_documentation('output')
